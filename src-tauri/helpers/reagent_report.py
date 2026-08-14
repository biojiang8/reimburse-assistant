#!/usr/bin/env python3
"""试剂耗材报销汇总工具。

输入：
  1. 供应商订单 Excel（含 项目名称/财务项目码/发票号/物品名称/物品规格/
     供应商编码/供应商名称/单位/单价/数量 列，表头可自动识别）
  2. 实物照片文件夹（文件名含货号者优先精确匹配，其余按明细顺序分配）
  3. 电子发票 PDF（可选；提供后按发票号核验金额并生成带电子签的重命名 PDF）

输出（到 --output-dir）：
  * <项目或订单名>-汇总.xlsx   订单明细表 + 发票汇总表
  * <项目或订单名>-实物证据.docx  按发票分组的 货号/名称/规格/数量 + 实物照片
  * pdf/ 目录                  带电子签的发票 PDF（提供发票时）

依赖：openpyxl、python-docx、PyMuPDF(fitz)、Pillow；可选复用同目录
add_invoice_signature.py 完成签字。
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import dataclass, field
from decimal import Decimal, InvalidOperation, ROUND_HALF_UP
from pathlib import Path
from typing import Dict, List, Optional, Sequence, Tuple

import fitz

try:
    import openpyxl
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter
except ImportError:  # pragma: no cover
    openpyxl = None  # type: ignore

try:
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
except ImportError:  # pragma: no cover
    docx = None  # type: ignore

from add_invoice_signature import (
    DEFAULT_SIGNATURE,
    InvoiceFields,
    add_signature,
    extract_invoice_fields,
    _money_text,
    _safe_filename_component,
)

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tif", ".tiff"}

# ---------------------------------------------------------------------------
# 数据模型
# ---------------------------------------------------------------------------


@dataclass
class OrderItem:
    row: int  # 订单表行号（用于报告）
    project_name: str
    project_code: str
    invoice_number: str
    name: str
    spec: str
    vendor_code: str
    vendor_name: str
    unit: str
    price: Decimal
    qty: Decimal
    photos: List[Path] = field(default_factory=list)


@dataclass
class InvoiceGroup:
    invoice_number: str
    seller: Optional[str] = None
    buyer: Optional[str] = None
    invoice_date: Optional[str] = None
    invoice_total: Optional[Decimal] = None
    order_total: Decimal = Decimal("0")
    verified: bool = False
    verify_note: str = ""
    signed_pdf: Optional[Path] = None
    source_pdf: Optional[Path] = None
    items: List[OrderItem] = field(default_factory=list)


# ---------------------------------------------------------------------------
# 订单表解析
# ---------------------------------------------------------------------------

ORDER_COLUMNS = [
    "项目名称",
    "财务项目码",
    "发票号",
    "物品名称",
    "物品规格",
    "供应商编码",
    "供应商名称",
    "单位",
    "单价",
    "数量",
]


def _cell_text(value: object) -> str:
    if value is None:
        return ""
    text = str(value).strip()
    # 数字单元格（如发票号）会被 Excel 转成科学计数法/浮点，尽量还原
    if isinstance(value, float) and value.is_integer() and len(text) <= 15:
        return str(int(value))
    return text


def _count_text(value: Decimal) -> str:
    """数量文本：整数不带小数位（1 → '1'），小数去掉末尾 0（1.5 → '1.5'）。"""
    quantized = value.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    text = format(quantized, "f")
    if "." in text:
        text = text.rstrip("0").rstrip(".")
    return text


def _parse_decimal(value: object, field_name: str, row: int) -> Decimal:
    text = _cell_text(value).replace(",", "").replace("￥", "").replace("¥", "")
    text = re.sub(r"[元个只箱包袋盒支瓶套]", "", text).strip()
    try:
        return Decimal(text).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    except (InvalidOperation, ValueError):
        raise ValueError(f"第 {row} 行 {field_name} 无法解析为数字：{value!r}")


def _find_header_row(ws) -> int:
    """定位表头行：同时包含 物品名称 和 发票号 列的行。"""
    for row_index, row in enumerate(ws.iter_rows(values_only=True), start=1):
        texts = {_cell_text(c) for c in row}
        if "物品名称" in texts and ("发票号" in texts or "发票号码" in texts):
            return row_index
    raise ValueError(
        "未找到订单表头（需要包含 物品名称、发票号 等列）。"
        "请使用供应商提供的标准订单表。"
    )


def _column_map(ws, header_row: int) -> Dict[str, int]:
    mapping: Dict[str, int] = {}
    for index, cell in enumerate(ws[header_row], start=1):
        text = _cell_text(cell.value)
        for column in ORDER_COLUMNS:
            if column in text and column not in mapping:
                mapping[column] = index
    missing = [column for column in ORDER_COLUMNS if column not in mapping]
    if missing:
        raise ValueError(f"订单表缺少列：{'、'.join(missing)}")
    return mapping


def parse_order_sheet(order_path: Path) -> List[OrderItem]:
    if openpyxl is None:
        raise RuntimeError("缺少 openpyxl，请先安装：pip3 install openpyxl")
    workbook = openpyxl.load_workbook(order_path, data_only=True)
    ws = workbook.worksheets[0]
    header_row = _find_header_row(ws)
    columns = _column_map(ws, header_row)

    items: List[OrderItem] = []
    for row_index in range(header_row + 1, ws.max_row + 1):
        values = list(ws[row_index])
        def col(column: str) -> object:
            index = columns[column]
            return values[index - 1].value if index <= len(values) else None

        name = _cell_text(col("物品名称"))
        invoice = _cell_text(col("发票号"))
        if not name or not invoice:
            continue  # 空行或底部备注行
        if "注意" in name or "请误删除" in name:
            continue
        price = _parse_decimal(col("单价"), "单价", row_index)
        qty = _parse_decimal(col("数量"), "数量", row_index)
        items.append(
            OrderItem(
                row=row_index,
                project_name=_cell_text(col("项目名称")),
                project_code=_cell_text(col("财务项目码")),
                invoice_number=invoice,
                name=name,
                spec=_cell_text(col("物品规格")),
                vendor_code=_cell_text(col("供应商编码")),
                vendor_name=_cell_text(col("供应商名称")),
                unit=_cell_text(col("单位")),
                price=price,
                qty=qty,
            )
        )
    if not items:
        raise ValueError(f"订单表中没有可用的明细行：{order_path}")
    return items


# ---------------------------------------------------------------------------
# 照片匹配
# ---------------------------------------------------------------------------


def _photo_stem(path: Path) -> str:
    stem = path.stem.strip()
    # 去掉常见的拍照序号后缀：_1 / -2 / (3) / 副本
    stem = re.sub(r"[_\-\s]*(\d{1,3}|副本|copy)$", "", stem, flags=re.IGNORECASE)
    return stem


def _normalize_code(value: str) -> str:
    return re.sub(r"[\s\-_（）()\[\]【】./\\]", "", value).upper()


def collect_photos(photos_dir: Optional[Path]) -> List[Path]:
    if photos_dir is None:
        return []
    if not photos_dir.is_dir():
        raise ValueError(f"照片文件夹不存在：{photos_dir}")
    photos = sorted(
        (p for p in photos_dir.iterdir() if p.suffix.lower() in IMAGE_EXTENSIONS),
        key=lambda p: p.name,
    )
    return photos


def assign_photos(items: List[OrderItem], photos: List[Path]) -> Tuple[List[Path], List[int]]:
    """先按货号精确匹配，剩余照片按明细顺序分配。

    返回 (未匹配照片, 缺照片的物品下标)。
    """
    used: set[int] = set()

    def code_of(item: OrderItem) -> str:
        return _normalize_code(item.vendor_code)

    # 第一轮：货号匹配（文件名包含货号，或货号包含文件名主体）
    for photo_index, photo in enumerate(photos):
        stem = _normalize_code(_photo_stem(photo))
        if not stem:
            continue
        for item_index, item in enumerate(items):
            code = code_of(item)
            if code and (stem == code or code in stem or stem in code):
                items[item_index].photos.append(photo)
                used.add(photo_index)
                break

    # 第二轮：剩余照片按物品顺序轮流分配
    remaining = [
        (index, photo) for index, photo in enumerate(photos) if index not in used
    ]
    pointer = 0
    for photo_index, photo in remaining:
        items[pointer % len(items)].photos.append(photo)
        used.add(photo_index)
        pointer += 1

    unmatched = [photo for index, photo in enumerate(photos) if index not in used]
    missing = [i for i, item in enumerate(items) if not item.photos]
    return unmatched, missing


# ---------------------------------------------------------------------------
# 发票提取与核验
# ---------------------------------------------------------------------------


def _find_invoice_date(pdf_path: Path) -> Optional[str]:
    with fitz.open(pdf_path) as document:
        if not document:
            return None
        text = re.sub(r"\s+", "", document[0].get_text("text"))
    # 先找与「开票日期」标签相邻的值
    for pattern in (
        r"开票日期[:：]?(\d{4}年\d{1,2}月\d{1,2}日)",
        r"开票日期[:：]?(\d{4}-\d{1,2}-\d{1,2})",
        r"开票日期[:：]?(\d{4}/\d{1,2}/\d{1,2})",
    ):
        match = re.search(pattern, text)
        if match:
            return match.group(1)
    # 扁平化文本层中标签与值可能分离：全文找第一个完整日期
    for pattern in (
        r"(\d{4}年\d{1,2}月\d{1,2}日)",
        r"(\d{4}-\d{1,2}-\d{1,2})",
        r"(\d{4}/\d{1,2}/\d{1,2})",
    ):
        match = re.search(pattern, text)
        if match:
            return match.group(1)
    return None


def _invoice_pdf_paths(values: Sequence[str]) -> List[Path]:
    paths: List[Path] = []
    for raw in values:
        path = Path(raw).expanduser()
        if path.is_dir():
            paths.extend(sorted(p for p in path.glob("*.pdf") if p.is_file()))
        elif path.is_file() and path.suffix.lower() == ".pdf":
            paths.append(path)
        else:
            raise FileNotFoundError(f"发票 PDF 不存在：{path}")
    return paths


def build_groups(
    items: List[OrderItem],
    invoice_pdfs: Sequence[Path],
    signature_path: Optional[Path],
    output_pdf_dir: Path,
    overwrite: bool,
) -> Tuple[List[InvoiceGroup], List[dict]]:
    """按发票号分组订单行，并用发票 PDF 核验 + 签字。返回 (分组, 错误列表)。"""
    by_number: Dict[str, List[OrderItem]] = {}
    for item in items:
        by_number.setdefault(item.invoice_number, []).append(item)

    groups: List[InvoiceGroup] = []
    for number, group_items in by_number.items():
        group = InvoiceGroup(
            invoice_number=number,
            order_total=sum((i.price * i.qty for i in group_items), Decimal("0")).quantize(
                Decimal("0.01"), rounding=ROUND_HALF_UP
            ),
            items=group_items,
        )
        groups.append(group)

    # 找与订单发票号匹配的发票 PDF
    errors: List[dict] = []
    for pdf in invoice_pdfs:
        try:
            fields = extract_invoice_fields(pdf)
        except Exception as exc:  # 单张失败不中断整体
            errors.append({"file": str(pdf), "error": str(exc)})
            continue
        group = next(
            (g for g in groups if g.invoice_number == fields.invoice_number), None
        )
        if group is None:
            errors.append(
                {
                    "file": str(pdf),
                    "error": f"发票号 {fields.invoice_number} 不在订单表中，已跳过",
                }
            )
            continue
        group.seller = fields.seller
        group.buyer = fields.buyer
        group.invoice_date = _find_invoice_date(pdf)
        group.invoice_total = fields.total
        group.source_pdf = pdf

        diff = (fields.total - group.order_total).copy_abs()
        if diff <= Decimal("0.01"):
            group.verified = True
            group.verify_note = "金额一致"
        else:
            group.verify_note = (
                f"金额不一致：发票 {_money_text(fields.total)} "
                f"≠ 订单 {_money_text(group.order_total)}"
            )

        # 生成带电子签的重命名 PDF
        try:
            output_pdf_dir.mkdir(parents=True, exist_ok=True)
            filename = (
                f"{_safe_filename_component(fields.seller)}-{_money_text(fields.total)}-"
                f"dzfp{fields.invoice_number}-已签字.pdf"
            )
            output = output_pdf_dir / filename
            sig = signature_path or (Path(__file__).resolve().parent / DEFAULT_SIGNATURE)
            if not sig.is_file():
                raise FileNotFoundError(f"电子签图片不存在：{sig}")
            add_signature(
                input_pdf=pdf,
                signature_path=sig,
                output_pdf=output,
                company_side="seller",
                amount_field="total",
                add_label=True,
                overwrite=overwrite,
            )
            group.signed_pdf = output
        except FileExistsError:
            group.signed_pdf = output  # type: ignore[assignment]
            group.verify_note += "（签字文件已存在）"
        except Exception as exc:
            group.verify_note += f"；签字失败：{exc}"

    return groups, errors


# ---------------------------------------------------------------------------
# Excel 汇总
# ---------------------------------------------------------------------------


def _styled_header(ws, row: int, count: int) -> None:
    fill = PatternFill("solid", fgColor="DDEBF7")
    border = Border(*[Side(style="thin", color="999999")] * 4)
    for column in range(1, count + 1):
        cell = ws.cell(row=row, column=column)
        cell.font = Font(bold=True)
        cell.fill = fill
        cell.border = border
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)


def build_excel(
    groups: List[InvoiceGroup],
    output_path: Path,
    project_name: str,
    unmatched_photos: List[Path],
) -> Path:
    if openpyxl is None:
        raise RuntimeError("缺少 openpyxl，请先安装：pip3 install openpyxl")
    workbook = openpyxl.Workbook()

    # ---- Sheet1 订单明细 ----
    ws = workbook.active
    ws.title = "订单明细"
    headers = [
        "序号",
        "项目名称",
        "财务项目码",
        "发票号",
        "物品名称",
        "物品规格",
        "供应商编码",
        "供应商名称",
        "单位",
        "单价",
        "数量",
        "金额",
        "照片文件",
        "发票核验",
    ]
    ws.append(headers)
    _styled_header(ws, 1, len(headers))
    border = Border(*[Side(style="thin", color="CCCCCC")] * 4)
    row_index = 2
    serial = 0
    for group in groups:
        for item in group.items:
            serial += 1
            photo_names = "；".join(p.name for p in item.photos) or "（缺照片）"
            ws.append(
                [
                    serial,
                    item.project_name,
                    item.project_code,
                    item.invoice_number,
                    item.name,
                    item.spec,
                    item.vendor_code,
                    item.vendor_name,
                    item.unit,
                    float(item.price),
                    float(item.qty),
                    float(item.price * item.qty),
                    photo_names,
                    group.verify_note,
                ]
            )
            for column in range(1, len(headers) + 1):
                cell = ws.cell(row=row_index, column=column)
                cell.border = border
                if column in (10, 11, 12):
                    cell.number_format = "0.00"
                if column == 13:
                    cell.alignment = Alignment(wrap_text=True, vertical="center")
            if not item.photos:
                ws.cell(row=row_index, column=13).font = Font(color="C00000")
            row_index += 1
    # 合计行
    ws.append(["合计", "", "", "", "", "", "", "", "", "", "", "", "", ""])
    total_row = row_index
    ws.merge_cells(start_row=total_row, start_column=1, end_row=total_row, end_column=11)
    total_cell = ws.cell(row=total_row, column=12)
    total_cell.value = float(sum((g.order_total for g in groups), Decimal("0")))
    total_cell.number_format = "0.00"
    for column in range(1, len(headers) + 1):
        ws.cell(row=total_row, column=column).font = Font(bold=True)
        ws.cell(row=total_row, column=column).border = border

    widths = [6, 26, 12, 22, 34, 26, 14, 24, 8, 10, 8, 10, 34, 22]
    for index, width in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(index)].width = width
    ws.freeze_panes = "A2"

    # ---- Sheet2 发票汇总 ----
    ws2 = workbook.create_sheet("发票汇总")
    headers2 = [
        "序号",
        "发票号",
        "销售方",
        "购买方",
        "开票日期",
        "发票价税合计",
        "订单金额合计",
        "核验状态",
        "签字文件",
        "发票来源",
    ]
    ws2.append(headers2)
    _styled_header(ws2, 1, len(headers2))
    row2 = 2
    for index, group in enumerate(groups, start=1):
        ws2.append(
            [
                index,
                group.invoice_number,
                group.seller or "（未提供发票）",
                group.buyer or "",
                group.invoice_date or "",
                float(group.invoice_total) if group.invoice_total is not None else "",
                float(group.order_total),
                "已核验一致" if group.verified else (group.verify_note or "未核验"),
                group.signed_pdf.name if group.signed_pdf else "",
                group.source_pdf.name if group.source_pdf else "",
            ]
        )
        status = ws2.cell(row=row2, column=8)
        status.font = Font(color="007A33" if group.verified else "C00000", bold=not group.verified)
        for column in range(1, len(headers2) + 1):
            ws2.cell(row=row2, column=column).border = border
        row2 += 1
    ws2.append([])
    ws2.append(
        [
            "",
            "",
            "",
            "",
            "合计",
            float(sum((g.invoice_total for g in groups), Decimal("0")))
            if any(g.invoice_total is not None for g in groups)
            else "",
            float(sum((g.order_total for g in groups), Decimal("0"))),
            "",
            "",
            "",
        ]
    )
    for column in range(1, len(headers2) + 1):
        cell = ws2.cell(row=row2, column=column)
        cell.font = Font(bold=True)
        cell.border = border
    if unmatched_photos:
        ws2.append([])
        ws2.append(["未匹配照片", "；".join(p.name for p in unmatched_photos)])
        ws2.merge_cells(start_row=row2 + 2, start_column=2, end_row=row2 + 2, end_column=10)
        ws2.cell(row=row2 + 2, column=2).alignment = Alignment(wrap_text=True)
    widths2 = [6, 22, 32, 26, 12, 14, 14, 26, 40, 30]
    for index, width in enumerate(widths2, start=1):
        ws2.column_dimensions[get_column_letter(index)].width = width
    ws2.freeze_panes = "A2"

    workbook.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# Word 实物证据
# ---------------------------------------------------------------------------


def _item_description(item: OrderItem) -> str:
    return (
        f"货号:{item.vendor_code} 名称:{item.name} 规格:{item.spec}"
        f"数量:{_count_text(item.qty)}{item.unit}"
    )


def build_word(groups: List[InvoiceGroup], output_path: Path, project_name: str) -> Path:
    if docx is None:
        raise RuntimeError("缺少 python-docx，请先安装：pip3 install python-docx")
    document = docx.Document()
    style = document.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(docx.oxml.ns.qn("w:eastAsia"), "宋体")

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{project_name} 试剂耗材实物证据")
    run.bold = True
    run.font.size = Pt(16)

    first = True
    for group in groups:
        if not first:
            document.add_page_break()
        first = False

        heading = document.add_paragraph()
        vendor = group.seller or (group.items[0].vendor_name if group.items else "供应商")
        run = heading.add_run(
            f"{vendor}　发票号：{group.invoice_number}"
            f"　价税合计：{_money_text(group.invoice_total) if group.invoice_total is not None else '—'} 元"
        )
        run.bold = True
        run.font.size = Pt(12)

        for item in group.items:
            description = document.add_paragraph()
            description.paragraph_format.space_before = Pt(8)
            run = description.add_run(_item_description(item))
            run.font.size = Pt(10.5)
            if not item.photos:
                note = document.add_paragraph()
                run = note.add_run("（缺实物照片）")
                run.font.color.rgb = docx.shared.RGBColor(0xC0, 0x00, 0x00)
                continue
            for photo in item.photos:
                picture = document.add_paragraph()
                picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = picture.add_run()
                run.add_picture(str(photo), width=Inches(5.2))
                caption = document.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = caption.add_run(photo.name)
                run.font.size = Pt(9)
                run.font.color.rgb = docx.shared.RGBColor(0x59, 0x59, 0x59)

    document.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# 主流程
# ---------------------------------------------------------------------------


def _project_title(items: List[OrderItem], order_path: Path) -> str:
    names = {i.project_name for i in items if i.project_name}
    if len(names) == 1:
        return next(iter(names))
    return order_path.stem


def run_report(
    order_path: Path,
    photos_dir: Optional[Path],
    invoice_pdfs: Sequence[Path],
    output_dir: Path,
    signature_path: Optional[Path],
    overwrite: bool,
) -> dict:
    items = parse_order_sheet(order_path)
    photos = collect_photos(photos_dir)
    unmatched_photos, missing_indexes = assign_photos(items, photos)

    output_dir.mkdir(parents=True, exist_ok=True)
    pdf_dir = output_dir / "pdf"
    groups, invoice_errors = build_groups(
        items, invoice_pdfs, signature_path, pdf_dir, overwrite
    )

    project_name = _project_title(items, order_path)
    excel_path = build_excel(groups, output_dir / f"{_safe_filename_component(project_name)}-汇总.xlsx", project_name, unmatched_photos)
    word_path = build_word(groups, output_dir / f"{_safe_filename_component(project_name)}-实物证据.docx", project_name)

    result = {
        "ok": True,
        "project_name": project_name,
        "order_file": str(order_path.resolve()),
        "photos_dir": str(photos_dir.resolve()) if photos_dir else None,
        "output_dir": str(output_dir.resolve()),
        "excel": str(excel_path.resolve()),
        "word": str(word_path.resolve()),
        "invoice_count": len(groups),
        "item_count": len(items),
        "photo_count": len(photos),
        "unmatched_photos": [str(p) for p in unmatched_photos],
        "missing_photo_items": [items[i].name for i in missing_indexes],
        "invoices": [
            {
                "invoice_number": g.invoice_number,
                "seller": g.seller,
                "buyer": g.buyer,
                "invoice_date": g.invoice_date,
                "invoice_total": _money_text(g.invoice_total) if g.invoice_total is not None else None,
                "order_total": _money_text(g.order_total),
                "verified": g.verified,
                "verify_note": g.verify_note,
                "signed_pdf": str(g.signed_pdf.resolve()) if g.signed_pdf else None,
                "items": [
                    {
                        "name": item.name,
                        "spec": item.spec,
                        "vendor_code": item.vendor_code,
                        "unit": item.unit,
                        "price": _money_text(item.price),
                        "qty": _money_text(item.qty),
                        "photos": [str(p) for p in item.photos],
                    }
                    for item in g.items
                ],
            }
            for g in groups
        ],
    }
    if invoice_errors:
        result["invoice_errors"] = invoice_errors
    return result


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="试剂耗材报销：订单表 + 实物照片 + 发票 → 汇总 Excel + Word 证据"
    )
    parser.add_argument("--order", type=Path, required=True, help="供应商订单 Excel")
    parser.add_argument("--photos-dir", type=Path, default=None, help="实物照片文件夹")
    parser.add_argument(
        "--invoices", nargs="+", default=[], help="电子发票 PDF 文件或目录（可选）"
    )
    parser.add_argument("--output-dir", type=Path, required=True, help="输出目录")
    parser.add_argument(
        "--signature", type=Path, default=None, help="电子签图片（默认使用内置签名）"
    )
    parser.add_argument("--overwrite", action="store_true", help="覆盖已存在的输出")
    parser.add_argument("--json", action="store_true", dest="json_output", help="输出 JSON")
    return parser


def main(argv: Optional[Sequence[str]] = None) -> int:
    args = build_parser().parse_args(argv)
    try:
        result = run_report(
            order_path=args.order,
            photos_dir=args.photos_dir,
            invoice_pdfs=_invoice_pdf_paths(args.invoices),
            output_dir=args.output_dir,
            signature_path=args.signature,
            overwrite=args.overwrite,
        )
        if args.json_output:
            print(json.dumps(result, ensure_ascii=False))
        else:
            print(f"项目：{result['project_name']}")
            print(f"发票 {result['invoice_count']} 张，明细 {result['item_count']} 行，照片 {result['photo_count']} 张")
            if result["unmatched_photos"]:
                print("未匹配照片：", "、".join(Path(p).name for p in result["unmatched_photos"]))
            if result["missing_photo_items"]:
                print("缺照片物品：", "、".join(result["missing_photo_items"]))
            for g in result["invoices"]:
                status = "已核验一致" if g["verified"] else (g["verify_note"] or "未核验")
                print(f"  {g['invoice_number']} 订单 {g['order_total']} 元 → {status}")
            print(f"Excel：{result['excel']}")
            print(f"Word：{result['word']}")
        return 0
    except (FileNotFoundError, ValueError, RuntimeError) as exc:
        if getattr(args, "json_output", False):
            print(json.dumps({"ok": False, "error": str(exc)}, ensure_ascii=False))
        else:
            print(f"error: {exc}", file=sys.stderr)
        return 2


if __name__ == "__main__":
    raise SystemExit(main())
