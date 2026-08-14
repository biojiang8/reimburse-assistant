#!/usr/bin/env python3
"""每张发票配套的"实物证据 Word"生成器（python-docx）。

内容：发票号码、公司（销售方）、税号、价税合计、开票日期、
以及拖入的实物照片（每张一行标题 + 图片）。

输入：--json-file 指向的记录 JSON（单个对象或含 records 字段的数组，
数组时取第一条）。字段：invoice_number / seller / seller_tax_id /
total / invoice_date / photos（路径数组）。
输出：--output 指向的 .docx 文件。
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from typing import Any, Dict, List


def _text(record: Dict[str, Any], key: str, default: str = "") -> str:
    value = record.get(key)
    return "" if value is None else str(value).strip()


def _record_from_data(data: Any) -> Dict[str, Any]:
    if isinstance(data, dict):
        return data
    if isinstance(data, list) and data:
        return data[0]
    return {}


def _style_run(run, bold: bool = False, size: int = 12) -> None:
    """统一字体：宋体 12pt（中文走 eastAsia 字体槽）。"""

    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = "宋体"
    run.font.size = Pt(size)
    run.bold = bold
    element = run._element
    rpr = element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), "宋体")


def _style_paragraph(paragraph) -> None:
    paragraph.paragraph_format.line_spacing = 1.5


def write_evidence(record: Dict[str, Any], output_path: Path) -> None:
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    document = docx.Document()

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _style_paragraph(title)
    _style_run(title.add_run("实物证据（报销附件）"), bold=True)

    info_lines = [
        ("发票号码", _text(record, "invoice_number") or "—"),
        ("公司（销售方）", _text(record, "seller") or "—"),
        ("税号", _text(record, "seller_tax_id") or "—"),
        ("价税合计（元）", _text(record, "total") or "—"),
        ("开票日期", _text(record, "invoice_date") or "—"),
    ]
    for label, value in info_lines:
        paragraph = document.add_paragraph()
        _style_paragraph(paragraph)
        paragraph.paragraph_format.space_after = Pt(4)
        _style_run(paragraph.add_run(f"{label}："), bold=True)
        _style_run(paragraph.add_run(value))

    photos = record.get("photos") or []
    photos = [str(p) for p in photos if p]
    if photos:
        document.add_paragraph()
        for index, photo in enumerate(photos, start=1):
            caption = document.add_paragraph()
            _style_paragraph(caption)
            caption.paragraph_format.space_before = Pt(10)
            _style_run(caption.add_run(f"实物照片 {index}"), bold=True)
            path = Path(photo)
            if path.is_file():
                try:
                    picture = document.add_picture(str(path), width=Cm(15))
                    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as exc:  # noqa: BLE001 — 图片损坏时保留文字占位
                    note = document.add_paragraph()
                    _style_paragraph(note)
                    _style_run(note.add_run(f"（图片无法嵌入：{exc}）"))
            else:
                note = document.add_paragraph()
                _style_paragraph(note)
                _style_run(note.add_run(f"（图片不存在：{path.name}）"))
    else:
        note = document.add_paragraph("（未提供实物照片）")
        _style_paragraph(note)
        _style_run(note.add_run("（未提供实物照片）"))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def main(argv: List[str]) -> int:
    parser = argparse.ArgumentParser(description="实物证据 Word 生成")
    # 由 reimburse-helper 的 evidence 子命令路由使用，脚本本身忽略
    parser.add_argument("--evidence-json", action="store_true", help=argparse.SUPPRESS)
    parser.add_argument("--json-file", required=True, help="记录 JSON 文件路径")
    parser.add_argument("--output", required=True, help="输出 .docx 路径")
    args = parser.parse_args(argv)

    data = json.loads(Path(args.json_file).read_text(encoding="utf-8"))
    record = _record_from_data(data)
    if not record:
        print(json.dumps({"ok": False, "error": "记录为空"}, ensure_ascii=False))
        return 2

    output_path = Path(args.output)
    write_evidence(record, output_path)
    print(
        json.dumps(
            {
                "ok": True,
                "word": str(output_path.resolve()),
                "photos": len(record.get("photos") or []),
            },
            ensure_ascii=False,
        )
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
